import json
from pathlib import Path
from typing import Dict, List, Optional

from PyPDF2 import PdfReader
from PyPDF2.errors import PdfReadError
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.exceptions import PackageNotFoundError

try:
    import win32com.client as win32  # type: ignore
    import pywintypes  # type: ignore
except ImportError:
    win32 = None
    pywintypes = None


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}
DEFAULT_INPUT_PATH = Path(r"C:\Users\Abhinav\Documents\Repo\Colab-proj-2\data\resumes")
DEFAULT_OUTPUT_PATH = Path("dataset/resume_text.jsonl")

RECOVERABLE_EXTRACTION_ERRORS = (
    ValueError,
    OSError,
    PdfReadError,
    PackageNotFoundError,
    KeyError,
    RuntimeError,
)
# if pywintypes is not None and hasattr(pywintypes, "com_error"):
#     RECOVERABLE_EXTRACTION_ERRORS = RECOVERABLE_EXTRACTION_ERRORS + (
#         pywintypes.com_error,  # type: ignore[attr-defined]
#     )


class WordAutomationClient:
    """Thin wrapper around Word COM automation to read legacy .doc files."""

    def __init__(self) -> None:
        if win32 is None:
            raise ImportError("pywin32 is required for .doc support on Windows")
        self._word = win32.Dispatch("Word.Application")
        self._word.Visible = False

    def close(self) -> None:
        if self._word is not None:
            self._word.Quit()
            self._word = None

    def __enter__(self) -> "WordAutomationClient":
        return self

    def __exit__(self, exc_type, exc, exc_tb) -> None:  # type: ignore[override]
        self.close()

    def extract_doc(self, file_path: Path) -> tuple[str, bool]:
        if self._word is None:
            raise RuntimeError("Word automation client is closed")
        document = self._word.Documents.Open(str(file_path))
        try:
            text = document.Content.Text
            contains_images = bool(document.InlineShapes.Count or document.Shapes.Count)
        finally:
            document.Close(False)
        return text.strip(), contains_images


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from a PDF file."""
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += (page.extract_text() or "") + "\n"
    return text.strip()


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from a DOCX file."""
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text.strip()


def extract_text(file_path: Path) -> str:
    """Extract text from a document file based on its extension."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    if suffix == ".docx":
        return extract_text_from_docx(file_path)
    raise ValueError(f"Unsupported file type: {file_path.suffix}")


def _xobject_dict_contains_images(x_objects) -> bool:
    if not x_objects:
        return False
    # Dereference x_objects if it's an indirect object
    try:
        x_objects = x_objects.get_object()
    except AttributeError:
        pass
    if not hasattr(x_objects, "values"):
        return False
    for obj in x_objects.values():
        try:
            x_obj = obj.get_object()
        except AttributeError:
            x_obj = obj
        subtype = x_obj.get("/Subtype")
        if subtype == "/Image":
            return True
        if subtype == "/Form":
            child_resources = x_obj.get("/Resources")
            child_x_objects = None
            if child_resources:
                # Dereference indirect objects
                try:
                    child_resources = child_resources.get_object()
                except AttributeError:
                    pass
                if hasattr(child_resources, "get"):
                    child_x_objects = child_resources.get("/XObject")
            if _xobject_dict_contains_images(child_x_objects):
                return True
    return False


def pdf_has_images(file_path: Path) -> bool:
    reader = PdfReader(file_path)
    for page in reader.pages:
        if hasattr(page, "images") and page.images:
            return True
        resources = page.get("/Resources")
        if not resources:
            continue
        # Dereference indirect objects
        try:
            resources = resources.get_object()
        except AttributeError:
            pass
        if not hasattr(resources, "get"):
            continue
        x_objects = resources.get("/XObject")
        if _xobject_dict_contains_images(x_objects):
            return True
    return False


def docx_has_images(file_path: Path) -> bool:
    doc = Document(file_path)
    return any(rel.reltype == RT.IMAGE for rel in doc.part.rels.values())


def has_images(file_path: Path) -> bool:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return pdf_has_images(file_path)
    if suffix == ".docx":
        return docx_has_images(file_path)
    raise ValueError(f"Unsupported file type: {file_path.suffix}")


def collect_documents(target: Path) -> List[Path]:
    if target.is_file():
        return [target]
    if target.is_dir():
        return sorted(
            [f for f in target.rglob("*") if f.suffix.lower() in SUPPORTED_EXTENSIONS]
        )
    raise FileNotFoundError(f"'{target}' is not a valid file or directory.")


def process_document(
    file_path: Path, word_client: Optional["WordAutomationClient"]
) -> Dict[str, str | bool]:
    suffix = file_path.suffix.lower()
    if suffix == ".doc":
        if word_client is None:
            raise RuntimeError(
                ".doc support requires Microsoft Word and pywin32; both appear unavailable."
            )
        resume_text, contains_images = word_client.extract_doc(file_path)
    else:
        resume_text = extract_text(file_path)
        contains_images = has_images(file_path)

    return {
        "filename": file_path.name,
        "filetype": suffix,
        "resume_text": resume_text,
        "has_images": contains_images,
    }


def main() -> None:
    # Update these paths to control which resumes are processed and where JSONL output lands.
    input_path = DEFAULT_INPUT_PATH
    output_path = DEFAULT_OUTPUT_PATH

    try:
        documents = collect_documents(input_path)
    except FileNotFoundError as exc:
        print(exc)
        return

    if not documents:
        print(f"No PDF/DOC/DOCX files found under '{input_path}'.")
        return

    needs_word = any(path.suffix.lower() == ".doc" for path in documents)
    word_client: Optional[WordAutomationClient] = None
    if needs_word:
        try:
            word_client = WordAutomationClient()
        except ImportError as exc:
            print(f"Cannot open .doc files: {exc}")
            return

    processed_count = 0
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        with output_path.open("w", encoding="utf-8") as fh:
            for doc_path in documents:
                try:
                    record = process_document(doc_path, word_client)
                except RECOVERABLE_EXTRACTION_ERRORS as exc:
                    print(f"Skipping {doc_path}: {exc}")
                    continue
                fh.write(json.dumps(record, ensure_ascii=False) + "\n")
                processed_count += 1
                print(f"Processed {doc_path.name}")
    finally:
        if word_client is not None:
            word_client.close()

    if processed_count:
        print(
            f"Wrote metadata for {processed_count} file(s) to '{output_path.as_posix()}'"
        )
    else:
        print("No files were successfully processed; see logs above for details.")


if __name__ == "__main__":
    main()